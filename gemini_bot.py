import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
import os
import time
import re
import webbrowser
import requests
from docx import Document
import google.generativeai as genai

class GeminiUniversalApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Gemini AI Assistant")
        self.root.geometry("550x850")
        
        # VERSION
        self.current_version = 1.0
        self.version_url = "https://raw.githubusercontent.com/Tamil-Venthan/Gemini-Assistant/refs/heads/main/version.txt"
        self.download_url = "https://github.com/Tamil-Venthan/Gemini-Assistant/releases"

        # VARIABLES
        self.input_path = tk.StringVar()
        self.model_var = tk.StringVar(value="models/gemini-flash-latest")
        self.temp_var = tk.DoubleVar(value=0.5)
        self.is_running = False
        self.stop_event = threading.Event()

        # LINKS
        self.linkedin_url = "https://www.linkedin.com/in/tamil-venthan4"
        self.github_url = "https://github.com/Tamil-Venthan"

        self.create_widgets()
        
        # Auto-check on startup
        threading.Thread(target=lambda: self.perform_update_check(silent=True), daemon=True).start()

    def create_widgets(self):
        # 1. Settings Frame
        config_frame = ttk.LabelFrame(self.root, text="Settings", padding=10)
        config_frame.pack(fill="x", padx=10, pady=5)

        ttk.Label(config_frame, text="Google API Key:").grid(row=0, column=0, sticky="w", pady=2)
        self.entry_api = ttk.Entry(config_frame, show="*", width=40)
        self.entry_api.grid(row=0, column=1, sticky="ew", padx=5, pady=2, columnspan=2)

        ttk.Label(config_frame, text="AI Model:").grid(row=1, column=0, sticky="w", pady=2)
        models = ["models/gemini-flash-latest", "models/gemini-2.0-flash", "models/gemini-1.5-flash", "models/gemini-1.5-pro"]
        self.combo_model = ttk.Combobox(config_frame, values=models, textvariable=self.model_var)
        self.combo_model.grid(row=1, column=1, sticky="w", padx=5, pady=2, columnspan=2)

        ttk.Label(config_frame, text="Creativity (0-1):").grid(row=2, column=0, sticky="w", pady=2)
        self.scale_temp = ttk.Scale(config_frame, from_=0.0, to=1.0, variable=self.temp_var, orient="horizontal")
        self.scale_temp.grid(row=2, column=1, sticky="ew", padx=5, pady=2)
        
        btn_info = ttk.Button(config_frame, text="?", width=3, command=self.show_temp_info)
        btn_info.grid(row=2, column=2, padx=5)

        # 2. System Instructions
        prompt_frame = ttk.LabelFrame(self.root, text="System Instructions", padding=10)
        prompt_frame.pack(fill="x", padx=10, pady=5)
        
        default_prompt = "You are a helpful expert. Answer clearly in bullet points. Use ONLY double asterisks (**) for bolding key terms. Do not use Markdown headers (###)."
        self.text_prompt = tk.Text(prompt_frame, height=4, width=50, font=("Arial", 9))
        self.text_prompt.insert("1.0", default_prompt)
        self.text_prompt.pack(fill="x")

        # 3. File Selection
        file_frame = ttk.LabelFrame(self.root, text="Select Question File (Word .docx)", padding=10)
        file_frame.pack(fill="x", padx=10, pady=5)

        self.entry_file = ttk.Entry(file_frame, textvariable=self.input_path, state="readonly")
        self.entry_file.pack(side="left", fill="x", expand=True, padx=(0, 5))
        
        btn_browse = ttk.Button(file_frame, text="Browse...", command=self.browse_file)
        btn_browse.pack(side="right")

        # 4. Action Area
        action_frame = ttk.Frame(self.root, padding=10)
        action_frame.pack(fill="x", padx=10)

        self.btn_est = ttk.Button(action_frame, text="Calculate Cost & Time", command=self.estimate_usage)
        self.btn_est.pack(fill="x", pady=5)

        btn_row = ttk.Frame(action_frame)
        btn_row.pack(fill="x", pady=5)

        self.btn_run = ttk.Button(btn_row, text="START / RESUME", command=self.start_thread)
        self.btn_run.pack(side="left", fill="x", expand=True, padx=(0, 5))

        self.btn_stop = ttk.Button(btn_row, text="STOP", command=self.stop_process, state="disabled")
        self.btn_stop.pack(side="right", fill="x", expand=True, padx=(5, 0))

        self.progress = ttk.Progressbar(action_frame, orient="horizontal", mode="determinate")
        self.progress.pack(fill="x", pady=5)

        # 5. Log
        log_frame = ttk.LabelFrame(self.root, text="Status Log", padding=10)
        log_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.log_area = scrolledtext.ScrolledText(log_frame, height=10, state='disabled', font=("Consolas", 9))
        self.log_area.pack(fill="both", expand=True)

        # 6. FOOTER
        footer_frame = ttk.Frame(self.root, padding=10)
        footer_frame.pack(fill="x", side="bottom")

        lbl_dev = ttk.Label(footer_frame, text="Contact Us", font=("Arial", 8, "italic"))
        lbl_dev.pack(pady=(0, 5))

        # MANUAL UPDATE BUTTON
        btn_update = ttk.Button(footer_frame, text="Check for Updates", command=lambda: self.perform_update_check(silent=False))
        btn_update.pack(pady=(0, 10))

        link_box = ttk.Frame(footer_frame)
        link_box.pack()

        btn_linkedin = ttk.Button(link_box, text="LinkedIn", width=12, command=lambda: self.open_link(self.linkedin_url))
        btn_linkedin.pack(side="left", padx=5)

        btn_github = ttk.Button(link_box, text="GitHub", width=12, command=lambda: self.open_link(self.github_url))
        btn_github.pack(side="left", padx=5)

    # SHARED UPDATE LOGIC
    def perform_update_check(self, silent=True):
        """
        Checks for updates.
        silent=True: No popup if up-to-date (used for auto-check).
        silent=False: Shows 'You are up to date' popup (used for button).
        """
        try:
            if not silent:
                self.log("Checking for updates...")
            
            # Simple request to get version number
            response = requests.get(self.version_url, timeout=5)
            
            if response.status_code == 200:
                remote_version = float(response.text.strip())
                
                if remote_version > self.current_version:
                    # UPDATE AVAILABLE
                    self.root.after(0, lambda: self.show_update_popup(remote_version))
                else:
                    # UP TO DATE
                    if not silent:
                        self.root.after(0, lambda: messagebox.showinfo("Up to Date", f"You have the latest version (v{self.current_version})."))
                        self.root.after(0, lambda: self.log("Check complete: Up to date."))

            else:
                if not silent:
                    self.root.after(0, lambda: messagebox.showerror("Error", "Could not reach update server."))

        except Exception as e:
            if not silent:
                self.root.after(0, lambda: messagebox.showerror("Error", f"Update check failed: {e}"))

    def show_update_popup(self, new_ver):
        msg = f"A new version (v{new_ver}) is available!\n\nWould you like to download it now?"
        if messagebox.askyesno("Update Available", msg):
            webbrowser.open_new(self.download_url)

    # STANDARD FUNCTIONS
    def open_link(self, url): webbrowser.open_new(url)
    def show_temp_info(self): messagebox.showinfo("Creativity", "0.0-0.3: Strict (Exams)\n0.5-0.7: Balanced\n0.8-1.0: Creative")
    
    def log(self, message):
        self.log_area.config(state='normal')
        self.log_area.insert(tk.END, message + "\n")
        self.log_area.see(tk.END)
        self.log_area.config(state='disabled')

    def browse_file(self):
        filename = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if filename: self.input_path.set(filename)

    def estimate_usage(self):
        input_file = self.input_path.get()
        if not input_file or not os.path.exists(input_file):
            messagebox.showerror("Error", "Please select a file first.")
            return
        try:
            doc = Document(input_file)
            questions = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            total_q = len(questions)
            if total_q == 0: return
            total_chars = sum(len(q) for q in questions)
            est_minutes = round((total_q * 6) / 60, 1)
            report = f"--- ESTIMATION REPORT ---\nTotal Questions: {total_q}\nEst. Tokens: ~{int(total_chars / 4)}\nESTIMATED TIME: {est_minutes} Minutes"
            messagebox.showinfo("Cost & Time Estimate", report)
        except Exception as e: messagebox.showerror("Error", f"Could not read file: {e}")

    def stop_process(self):
        if self.is_running:
            self.stop_event.set()
            self.log("!!! STOPPING REQUESTED... !!!")
            self.btn_stop.config(state="disabled")

    def start_thread(self):
        if self.is_running: return
        api_key = self.entry_api.get().strip()
        input_file = self.input_path.get()
        if not api_key: messagebox.showerror("Error", "Enter Google API Key."); return
        if not input_file: messagebox.showerror("Error", "Select input file."); return

        self.is_running = True
        self.stop_event.clear()
        self.btn_run.config(state="disabled")
        self.btn_stop.config(state="normal")
        threading.Thread(target=self.run_process, args=(api_key, input_file), daemon=True).start()

    def run_process(self, api_key, input_file):
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name=self.model_var.get(), system_instruction=self.text_prompt.get("1.0", tk.END).strip())
            gen_config = genai.GenerationConfig(temperature=self.temp_var.get())
            doc_in = Document(input_file)
            questions = [p.text.strip() for p in doc_in.paragraphs if p.text.strip()]
            total_items = len(questions)
            if total_items == 0: self.log("Error: No text found."); self.reset_ui(); return

            output_filename = input_file.replace(".docx", "_Gemini_Answers.docx")
            processed_count = 0
            if os.path.exists(output_filename):
                self.log(f"Resuming file: {os.path.basename(output_filename)}")
                doc_out = Document(output_filename)
                for p in doc_out.paragraphs:
                    if p.text.startswith("Item ") and ":" in p.text: processed_count += 1
            else:
                self.log("Starting fresh...")
                doc_out = Document()
                doc_out.add_heading("Gemini Answers", 0)

            for i, question in enumerate(questions):
                if i < processed_count: continue
                if self.stop_event.is_set(): self.log("--- STOPPED ---"); break
                self.log(f"Processing Item {i+1}/{total_items}...")
                self.progress['value'] = ((i + 1) / total_items) * 100
                self.root.update_idletasks()
                try:
                    response = model.generate_content(question, generation_config=gen_config)
                    answer = response.text.strip()
                except Exception as e: answer = f"Error: {e}"; self.log(f"API Error: {e}")

                doc_out.add_heading(f"Item {i+1}: {question[:50]}...", level=2)
                self.add_strict_clean_markdown(doc_out, answer)
                doc_out.add_paragraph("_" * 40)
                if (i+1) % 5 == 0: doc_out.save(output_filename)
                time.sleep(5)

            doc_out.save(output_filename)
            if not self.stop_event.is_set():
                self.log("DONE!"); messagebox.showinfo("Success", "Process Completed!")
            else:
                self.log("Stopped."); messagebox.showinfo("Stopped", "Process Stopped.")
        except Exception as e: self.log(f"CRITICAL ERROR: {e}"); messagebox.showerror("Error", str(e))
        finally: self.reset_ui()

    def add_strict_clean_markdown(self, doc, text):
        for line in text.split('\n'):
            line = line.replace('***', '**').replace('__', '**').strip()
            if not line or line in ['*', '-', '_', '***', '---']: continue
            if line.startswith('#'):
                doc.add_paragraph().add_run(re.sub(r'^#+\s*', '', line)).bold = True
            elif line.startswith('* ') or line.startswith('- '):
                self.format_regex_bold(doc.add_paragraph(style='List Bullet'), line[2:].strip())
            else:
                self.format_regex_bold(doc.add_paragraph(), line)

    def format_regex_bold(self, paragraph, text):
        for i, part in enumerate(re.split(r'\*\*(.*?)\*\*', text)):
            if part: paragraph.add_run(part).bold = (i % 2 == 1)

    def reset_ui(self):
        self.is_running = False; self.stop_event.clear()
        self.btn_run.config(state="normal"); self.btn_stop.config(state="disabled"); self.progress['value'] = 0

if __name__ == "__main__":
    root = tk.Tk()
    app = GeminiUniversalApp(root)
    root.mainloop()